import os
import re
from io import BytesIO
from pathlib import Path

from bs4 import BeautifulSoup, NavigableString
from django.conf import settings
from django.db import models
from django.template import Context, Engine
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from html4docx import HtmlToDocx
from xhtml2pdf import pisa
from xhtml2pdf.files import pisaFileObject

from common.services.numbers import number_to_words_uz
from contracts.models import Contract

_TEMPLATE_ENGINE = Engine(builtins=["contracts.templatetags.contracts_extras"])

_orig_get_named_file = pisaFileObject.getNamedFile


def _get_named_file(self):
    uri = str(self.uri or "")
    if os.path.isfile(uri):
        return uri
    return _orig_get_named_file(self)


pisaFileObject.getNamedFile = _get_named_file

PLACEHOLDER_RE = re.compile(r"{{\s*([\w.]+)\s*}}")

_FONT_DIR = Path(settings.BASE_DIR) / "static" / "fonts"

_FONT_FACES = {
    "regular": (_FONT_DIR / "DejaVuSans.ttf").as_posix(),
    "bold": (_FONT_DIR / "DejaVuSans-Bold.ttf").as_posix(),
    "italic": (_FONT_DIR / "DejaVuSans-Oblique.ttf").as_posix(),
    "bold_italic": (_FONT_DIR / "DejaVuSans-BoldOblique.ttf").as_posix(),
}

_PDF_FONT_ALIASES = (
    "DejaVuSans",
    "DejaVu Sans",
    "Arial",
    "Helvetica",
    "Calibri",
    "Cambria",
    "Verdana",
    "Tahoma",
    "Georgia",
    "Segoe UI",
    "Times New Roman",
    "Times",
    "sans-serif",
    "serif",
)

_PDF_PAGE_CSS = '@page { size: A4; margin: 2cm; } body { font-family: "DejaVuSans"; font-size: 12pt; }'

_HEAD_TAG_RE = re.compile(r"<head[^>]*>", re.IGNORECASE)

_CSS_COMMENT_RE = re.compile(r"/\*.*?\*/", re.DOTALL)
_CSS_AT_BLOCK_RE = re.compile(r"@[\w-]+[^{}]*\{(?:[^{}]*\{[^{}]*\})*[^{}]*\}", re.DOTALL)
_CSS_AT_LINE_RE = re.compile(r"@[^{};]*;")
_CSS_RULE_RE = re.compile(r"([^{}]+)\{([^{}]*)\}")
_IMPORTANT_RE = re.compile(r"\s*!important\s*", re.IGNORECASE)

_INHERITED_PROPS = frozenset(
    {
        "color",
        "font-family",
        "font-size",
        "font-style",
        "font-weight",
        "letter-spacing",
        "line-height",
        "text-align",
        "text-indent",
        "text-transform",
    }
)

_UA_DEFAULT_RULES = (
    ("h1", {"font-size": "24pt", "font-weight": "bold", "color": "#000000"}),
    ("h2", {"font-size": "18pt", "font-weight": "bold", "color": "#000000"}),
    ("h3", {"font-size": "14pt", "font-weight": "bold", "color": "#000000"}),
    ("h4", {"font-size": "12pt", "font-weight": "bold", "color": "#000000"}),
    ("h5", {"font-size": "10pt", "font-weight": "bold", "color": "#000000"}),
    ("h6", {"font-size": "8pt", "font-weight": "bold", "color": "#000000"}),
    ("th", {"text-align": "center"}),
    ("center", {"text-align": "center"}),
)

_GENERIC_FONT_MAP = {
    "sans-serif": "Arial",
    "serif": "Times New Roman",
    "monospace": "Courier New",
    "dejavusans": "DejaVu Sans",
}

_CSS_SIZE_RE = re.compile(r"([\d.]+)\s*(px|pt|em|rem|%)?")

_DOCX_DEFAULT_FONT = "Arial"

_TBLPR_SEQUENCE = (
    "tblStyle",
    "tblpPr",
    "tblOverlap",
    "bidiVisual",
    "tblStyleRowBandSize",
    "tblStyleColBandSize",
    "tblW",
    "jc",
    "tblCellSpacing",
    "tblInd",
    "tblBorders",
    "shd",
    "tblLayout",
    "tblCellMar",
    "tblLook",
)

_DEFAULT_CELL_MARGINS_DXA = {"top": 0, "left": 108, "bottom": 0, "right": 108}


def extract_placeholders(html: str) -> list[str]:
    return sorted(set(PLACEHOLDER_RE.findall(html)))


class BlankNoneWrapper:
    def __init__(self, obj):
        self._obj = obj

    def __getattr__(self, name):
        value = getattr(self._obj, name)
        if value is None:
            return ""
        if isinstance(value, models.Model):
            return BlankNoneWrapper(value)
        return value

    def __str__(self):
        return str(self._obj)


def _wrap(value):
    if isinstance(value, models.Model):
        return BlankNoneWrapper(value)
    if value is None:
        return ""
    return value


def build_context(contract: Contract) -> dict:
    ctx = dict(contract.data or {})
    for key, value in list(ctx.items()):
        words = number_to_words_uz(value)
        if words:
            ctx.setdefault(f"{key}_sozda", words)
    booking = contract.booking
    if booking is not None:
        ctx.setdefault("booking", booking)
        ctx.setdefault("client", booking.client)
        ctx.setdefault("home", booking.home)
        ctx.setdefault("company", booking.company)
        ctx.setdefault("total_price", booking.total_price)
        ctx.setdefault("total_price_sozda", booking.total_price_inword)
        ctx.setdefault("guarantee_percent", booking.guarantee_percent)
        ctx.setdefault("price_per_m2", booking.price_per_m2)
    ctx.setdefault("contract", contract)
    return {key: _wrap(value) for key, value in ctx.items()}


def render_contract_html(contract: Contract) -> str:
    source = contract.template.read_html()
    return _TEMPLATE_ENGINE.from_string(source).render(Context(build_context(contract)))


def _font_face_css(family: str) -> str:
    return (
        f'@font-face {{ font-family: "{family}"; src: url("{_FONT_FACES["regular"]}"); }}'
        f'@font-face {{ font-family: "{family}"; src: url("{_FONT_FACES["bold"]}"); font-weight: bold; }}'
        f'@font-face {{ font-family: "{family}"; src: url("{_FONT_FACES["italic"]}"); font-style: italic; }}'
        f'@font-face {{ font-family: "{family}"; src: url("{_FONT_FACES["bold_italic"]}");'
        f" font-weight: bold; font-style: italic; }}"
    )


def _build_pdf_base_css(html: str) -> str:
    html_lower = html.lower()
    faces = [
        _font_face_css(family) for family in _PDF_FONT_ALIASES if family == "DejaVuSans" or family.lower() in html_lower
    ]
    return "<style>" + "".join(faces) + _PDF_PAGE_CSS + "</style>"


def _pdf_link_callback(uri: str, rel: str) -> str:
    if uri.startswith(("data:", "http://", "https://")) or os.path.isfile(uri):
        return uri
    relative = uri.lstrip("/")
    targets = (
        (settings.MEDIA_URL.lstrip("/"), (settings.MEDIA_ROOT,)),
        (
            settings.STATIC_URL.lstrip("/"),
            (getattr(settings, "STATIC_ROOT", None), Path(settings.BASE_DIR) / "static"),
        ),
    )
    for prefix, roots in targets:
        if prefix and relative.startswith(prefix):
            rest = relative[len(prefix) :].lstrip("/")
            for root in roots:
                if root:
                    path = os.path.join(str(root), rest)
                    if os.path.isfile(path):
                        return path
    return uri


def html_to_pdf(html: str) -> bytes:
    base_css = _build_pdf_base_css(html)
    match = _HEAD_TAG_RE.search(html)
    insert_at = match.end() if match else 0
    html = html[:insert_at] + base_css + html[insert_at:]
    buffer = BytesIO()
    result = pisa.CreatePDF(src=html, dest=buffer, encoding="utf-8", link_callback=_pdf_link_callback)
    if result.err:
        raise ValueError("PDF generatsiya qilishda xatolik yuz berdi.")
    return buffer.getvalue()


def _parse_declarations(text: str) -> dict:
    decls = {}
    for part in text.split(";"):
        if ":" not in part:
            continue
        prop, value = part.split(":", 1)
        prop, value = prop.strip().lower(), value.strip()
        if prop and value:
            decls[prop] = value
    return decls


def _specificity(selector: str) -> tuple[int, int, int]:
    ids = selector.count("#")
    classes = selector.count(".") + selector.count("[") + selector.count(":")
    tags = len(re.findall(r"(?:^|[>+~\s])\s*([a-zA-Z][\w-]*)", selector))
    return ids, classes, tags


def _collect_css_rules(soup: BeautifulSoup) -> list:
    css = "\n".join(tag.get_text() for tag in soup.find_all("style"))
    css = _CSS_COMMENT_RE.sub("", css)
    css = _CSS_AT_BLOCK_RE.sub("", css)
    css = _CSS_AT_LINE_RE.sub("", css)
    rules = []
    for order, match in enumerate(_CSS_RULE_RE.finditer(css)):
        decls = _parse_declarations(match.group(2))
        if not decls:
            continue
        for selector in match.group(1).split(","):
            selector = selector.strip()
            if selector:
                rules.append((selector, _specificity(selector), order, decls))
    return rules


def _match_rules(soup: BeautifulSoup, rules: list) -> dict:
    matched: dict[int, list] = {}
    for selector, decls in _UA_DEFAULT_RULES:
        for el in soup.select(selector):
            matched.setdefault(id(el), []).append((((0, 0, 0), -1), decls))
    for selector, spec, order, decls in rules:
        try:
            selected = soup.select(selector)
        except Exception:
            continue
        for el in selected:
            matched.setdefault(id(el), []).append(((spec, order), decls))
    return matched


def _serialize_styles(styles: dict) -> str:
    return "; ".join(f"{prop}: {_IMPORTANT_RE.sub('', value)}" for prop, value in styles.items())


def _stamp_element(el, matched: dict, inherited: dict) -> dict:
    own = {}
    for _, decls in sorted(matched.get(id(el), []), key=lambda item: item[0]):
        own.update(decls)
    own.update(_parse_declarations(el.get("style", "")))
    computed = dict(inherited)
    computed.update(own)
    if computed:
        el["style"] = _serialize_styles(computed)
    next_inherited = {prop: value for prop, value in computed.items() if prop in _INHERITED_PROPS}
    if el.name in ("td", "th") and next_inherited:
        for child in list(el.children):
            if isinstance(child, NavigableString) and child.strip():
                span = BeautifulSoup("", "html.parser").new_tag("span")
                span["style"] = _serialize_styles(next_inherited)
                child.wrap(span)
    for child in el.find_all(recursive=False):
        _stamp_element(child, matched, next_inherited)
    return computed


def _inline_styles(soup: BeautifulSoup) -> dict:
    rules = _collect_css_rules(soup)
    matched = _match_rules(soup, rules)
    for tag in soup.find_all(["style", "script", "link", "meta", "title"]):
        tag.decompose()
    root = soup.body or soup
    seed = {"font-family": _DOCX_DEFAULT_FONT}
    if soup.body is not None:
        body_style = _stamp_element(soup.body, matched, seed)
        soup.body.attrs.pop("style", None)
    else:
        body_style = dict(seed)
        for child in root.find_all(recursive=False):
            _stamp_element(child, matched, seed)
    return body_style


def _css_size_to_pt(value: str | None):
    if not value:
        return None
    match = _CSS_SIZE_RE.match(_IMPORTANT_RE.sub("", value).strip())
    if not match:
        return None
    number = float(match.group(1))
    unit = match.group(2) or "px"
    factors = {"px": 0.75, "pt": 1.0, "em": 12.0, "rem": 12.0, "%": 0.12}
    factor = factors.get(unit)
    return Pt(round(number * factor, 1)) if factor else None


def _first_font_family(value: str | None):
    if not value:
        return None
    name = _IMPORTANT_RE.sub("", value).split(",")[0].strip().strip("'\"")
    if not name:
        return None
    return _GENERIC_FONT_MAP.get(name.lower(), name)


def _apply_document_defaults(document, body_style: dict) -> None:
    section = document.sections[0]
    section.page_width, section.page_height = Cm(21.0), Cm(29.7)
    section.top_margin = section.bottom_margin = Cm(2.0)
    section.left_margin = section.right_margin = Cm(2.0)

    font = document.styles["Normal"].font
    font.size = _css_size_to_pt(body_style.get("font-size")) or Pt(12)
    family = _first_font_family(body_style.get("font-family"))
    if family:
        font.name = family


def _local_name(element) -> str:
    return element.tag.rsplit("}", 1)[-1]


def _tblpr_insert(tbl_pr, element) -> None:
    pos = _TBLPR_SEQUENCE.index(_local_name(element))
    for child in tbl_pr:
        local = _local_name(child)
        if local in _TBLPR_SEQUENCE and _TBLPR_SEQUENCE.index(local) > pos:
            child.addprevious(element)
            return
    tbl_pr.append(element)


def _expand_box_values(value: str) -> dict:
    parts = _IMPORTANT_RE.sub("", value).split()
    if not parts:
        return {}
    if len(parts) == 1:
        parts = parts * 4
    elif len(parts) == 2:
        parts = [parts[0], parts[1], parts[0], parts[1]]
    elif len(parts) == 3:
        parts = [parts[0], parts[1], parts[2], parts[1]]
    return {"top": parts[0], "right": parts[1], "bottom": parts[2], "left": parts[3]}


def _set_table_width(table, width_value: str) -> None:
    value = _IMPORTANT_RE.sub("", width_value).strip()
    tbl_w = OxmlElement("w:tblW")
    if value.endswith("%"):
        try:
            pct = float(value[:-1])
        except ValueError:
            return
        tbl_w.set(qn("w:type"), "pct")
        tbl_w.set(qn("w:w"), str(int(pct * 50)))
    else:
        size = _css_size_to_pt(value)
        if size is None:
            return
        tbl_w.set(qn("w:type"), "dxa")
        tbl_w.set(qn("w:w"), str(int(size.pt * 20)))
    tbl_pr = table._tbl.tblPr
    for old in tbl_pr.findall(qn("w:tblW")):
        tbl_pr.remove(old)
    _tblpr_insert(tbl_pr, tbl_w)


def _cell_padding(table_soup) -> dict | None:
    for cell in table_soup.find_all(["td", "th"]):
        if cell.find_parent("table") is not table_soup:
            continue
        styles = _parse_declarations(cell.get("style", ""))
        box = _expand_box_values(styles["padding"]) if "padding" in styles else {}
        for side in ("top", "right", "bottom", "left"):
            if f"padding-{side}" in styles:
                box[side] = styles[f"padding-{side}"]
        if box:
            return box
    return None


def _set_cell_margins(table, box: dict) -> None:
    margins = OxmlElement("w:tblCellMar")
    for side in ("top", "left", "bottom", "right"):
        size = _css_size_to_pt(box.get(side))
        dxa = int(size.pt * 20) if size is not None else _DEFAULT_CELL_MARGINS_DXA[side]
        element = OxmlElement(f"w:{side}")
        element.set(qn("w:w"), str(max(0, dxa)))
        element.set(qn("w:type"), "dxa")
        margins.append(element)
    tbl_pr = table._tbl.tblPr
    for old in tbl_pr.findall(qn("w:tblCellMar")):
        tbl_pr.remove(old)
    _tblpr_insert(tbl_pr, margins)


def _postprocess_tables(document, soup: BeautifulSoup) -> None:
    soup_tables = [t for t in soup.find_all("table") if t.find_parent("table") is None]
    for table, table_soup in zip(document.tables, soup_tables, strict=False):
        styles = _parse_declarations(str(table_soup.get("style") or ""))
        if "width" in styles:
            _set_table_width(table, styles["width"])
        padding = _cell_padding(table_soup)
        if padding:
            _set_cell_margins(table, padding)


def _fix_border_sizes(document) -> None:
    for border in document.element.xpath(".//w:tcBorders/*"):
        size = border.get(qn("w:sz"))
        if size is None:
            continue
        try:
            border.set(qn("w:sz"), str(max(2, round(float(size)))))
        except ValueError:
            continue


def _is_visually_empty(paragraph) -> bool:
    if paragraph.text.strip():
        return False
    return not paragraph._p.xpath(".//w:br | .//w:drawing | .//w:pict")


def _remove_empty_paragraphs(document) -> None:
    paragraphs = document.paragraphs
    remaining = len(paragraphs)
    for paragraph in paragraphs:
        if remaining <= 1 or not _is_visually_empty(paragraph):
            continue
        p = paragraph._p
        prev_el, next_el = p.getprevious(), p.getnext()
        prev_is_tbl = prev_el is not None and _local_name(prev_el) == "tbl"
        next_is_tbl = next_el is not None and _local_name(next_el) == "tbl"
        next_is_end = next_el is None or _local_name(next_el) == "sectPr"
        if prev_is_tbl and (next_is_tbl or next_is_end):
            continue
        p.getparent().remove(p)
        remaining -= 1


def _apply_paragraph_spacing(document, soup: BeautifulSoup, body_style: dict) -> None:
    first_p = soup.find("p")
    if first_p is None:
        return
    styles = _parse_declarations(str(first_p.get("style") or ""))
    bottom = styles.get("margin-bottom")
    if bottom is None and "margin" in styles:
        bottom = _expand_box_values(styles["margin"]).get("bottom")
    if bottom is None:
        space = _css_size_to_pt(body_style.get("font-size")) or Pt(12)
    else:
        space = _css_size_to_pt(bottom) or Pt(0)
    for paragraph in document.paragraphs:
        if paragraph.paragraph_format.space_after is None:
            paragraph.paragraph_format.space_after = space


def html_to_docx(html: str) -> bytes:
    soup = BeautifulSoup(html, "html.parser")
    body_style = _inline_styles(soup)
    document = HtmlToDocx().parse_html_string(str(soup))
    _apply_document_defaults(document, body_style)
    _postprocess_tables(document, soup)
    _fix_border_sizes(document)
    _remove_empty_paragraphs(document)
    _apply_paragraph_spacing(document, soup, body_style)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
